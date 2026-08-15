#!/usr/bin/env python3
"""Regression tests for DOCX fidelity issues found in real conversions."""

from __future__ import annotations

import base64
import contextlib
import importlib.util
import io
import json
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document

SKILL_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = SKILL_DIR / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

TINY_PNG = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)


def load_module(name: str):
    spec = importlib.util.spec_from_file_location(name, SCRIPTS_DIR / f"{name}.py")
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def write_tiny_png(path: Path) -> None:
    path.write_bytes(TINY_PNG)


def write_template_files(root: Path, relative_paths: list[str]) -> None:
    """写入用于模板签名测试的最小占位文件。

    Args:
        root (Path): 临时模板根目录。
        relative_paths (list[str]): 需要创建的相对路径列表。
    """
    for relative_path in relative_paths:
        path = root / relative_path
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text("placeholder\n", encoding="utf-8")


def basic_metadata_yaml(*, title_cn: str = "测试题目", extra: str = "") -> str:
    return (
        "report_style: 1\n"
        f"thesis_title_cn: {title_cn}\n"
        "thesis_title_en: Test Title\n"
        "college: 测试学院\n"
        "major: 测试专业\n"
        "name: 测试姓名\n"
        "student_id: 20260001\n"
        "mentor: 张老师\n"
        "class_name: 测试班级\n"
        "date: 2026年6月\n"
        f"{extra}"
    )


def rewrite_docx_xml(
    docx_path: Path, replacements: dict[str, str], additions: dict[str, str] | None = None
) -> None:
    original = docx_path.read_bytes()
    with tempfile.TemporaryDirectory() as tmp:
        original_zip = Path(tmp) / "original.docx"
        rewritten_zip = Path(tmp) / "rewritten.docx"
        original_zip.write_bytes(original)
        with (
            zipfile.ZipFile(original_zip, "r") as source,
            zipfile.ZipFile(rewritten_zip, "w") as target,
        ):
            for info in source.infolist():
                data = source.read(info.filename)
                if info.filename in replacements:
                    data = replacements[info.filename].encode("utf-8")
                target.writestr(info, data)
            for filename, text in (additions or {}).items():
                target.writestr(filename, text.encode("utf-8"))
        docx_path.write_bytes(rewritten_zip.read_bytes())


def test_import_docx_preserves_superscript_runs():
    import_docx = load_module("import_docx")
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("引用")
    ref = paragraph.add_run("1")
    ref.font.superscript = True

    runs = import_docx.run_payload(paragraph)
    assert runs == [
        {
            "index": 1,
            "text": "引用",
            "bold": False,
            "italic": False,
            "superscript": False,
            "subscript": False,
            "font_size_pt": None,
        },
        {
            "index": 2,
            "text": "1",
            "bold": False,
            "italic": False,
            "superscript": True,
            "subscript": False,
            "font_size_pt": None,
        },
    ]


def test_render_chapters_preserves_superscript_and_heading_levels():
    render_chapters = load_module("render_chapters")
    assert (
        render_chapters.block_to_latex({"semantic_role": "heading", "level": 2, "text": "二级标题"})
        == "\\section{二级标题}"
    )
    assert (
        render_chapters.block_to_latex(
            {
                "source_type": "paragraph",
                "text": "引用1",
                "runs": [
                    {"text": "引用", "superscript": False},
                    {"text": "1", "superscript": True},
                ],
            }
        )
        == "引用\\textsuperscript{1}\n"
    )
    assert (
        render_chapters.block_to_latex(
            {
                "source_type": "paragraph",
                "runs": [
                    {"text": '"产品', "superscript": False},
                    {"text": '差异化"', "superscript": False},
                ],
            }
        )
        == "``产品差异化''\n"
    )


def test_render_chapters_outputs_labels_and_reference_rewrites():
    render_chapters = load_module("render_chapters")
    figure_latex = render_chapters.block_to_latex(
        {
            "source_type": "image",
            "asset_output": "Images/word_media/image3.png",
            "caption": "年龄分布图",
            "label": "fig:age-distribution",
        }
    )
    assert r"\caption{年龄分布图}" in figure_latex
    assert r"\label{fig:age-distribution}" in figure_latex

    table_latex = render_chapters.block_to_latex(
        {
            "source_type": "table",
            "caption": "样本信息",
            "label": "tab:sample-info",
            "table": {"rows": [["指标", "值"], ["样本", "205"]]},
        }
    )
    assert r"\caption{样本信息}" in table_latex
    assert r"\label{tab:sample-info}" in table_latex

    paragraph_latex = render_chapters.block_to_latex(
        {
            "source_type": "paragraph",
            "text": "具体可见图2.1和表 1.2。",
            "reference_rewrites": [
                {
                    "source_text": "图2.1",
                    "target_kind": "figure",
                    "target_label": "fig:age-distribution",
                },
                {
                    "source_text": "表 1.2",
                    "target_kind": "table",
                    "target_label": "tab:sample-info",
                },
            ],
        }
    )
    assert "图2.1" not in paragraph_latex
    assert "表 1.2" not in paragraph_latex
    assert r"图~\ref{fig:age-distribution}" in paragraph_latex
    assert r"表~\ref{tab:sample-info}" in paragraph_latex

    overlapping_latex = render_chapters.block_to_latex(
        {
            "source_type": "paragraph",
            "text": "见图2.1、图2.10。",
            "reference_rewrites": [
                {
                    "source_text": "图2.1",
                    "target_kind": "figure",
                    "target_label": "fig:short",
                },
                {
                    "source_text": "图2.10",
                    "target_kind": "figure",
                    "target_label": "fig:long",
                },
            ],
        }
    )
    assert r"图~\ref{fig:short}、图~\ref{fig:long}" in overlapping_latex
    assert r"\ref{fig:short}0" not in overlapping_latex


def test_check_template_requires_new_fonts_directory_layout():
    check_template = load_module("check_template")
    base_signature = [
        "main.tex",
        "zufe.cls",
        "Reference.bib",
        "chapters/basicinfo.tex",
        "chapters/mainbody.tex",
        "misc/cover.tex",
        "misc/abstract.tex",
        "misc/originality.tex",
        "misc/reference.tex",
        "InitFile/schoolLogo.png",
    ]

    with tempfile.TemporaryDirectory() as tmp:
        old_root = Path(tmp) / "old"
        write_template_files(
            old_root,
            base_signature + ["simhei.ttf", "stsong.ttf", "stkaiti.ttf"],
        )
        old_result = check_template.check_template(old_root)
        assert old_result["status"] == "blocked"
        assert old_result["missing"] == [
            "fonts/simhei.ttf",
            "fonts/stsong.ttf",
            "fonts/stkaiti.ttf",
        ]

        new_root = Path(tmp) / "new"
        write_template_files(
            new_root,
            base_signature + ["fonts/simhei.ttf", "fonts/stsong.ttf", "fonts/stkaiti.ttf"],
        )
        new_result = check_template.check_template(new_root)
        assert new_result["status"] == "passed"


def test_check_template_guides_template_download_fallbacks():
    check_template = load_module("check_template")
    with tempfile.TemporaryDirectory() as tmp:
        result = check_template.check_template(Path(tmp))
        next_steps = "\n".join(result["next_steps"])
        assert "https://github.com/sqsssq/ZUFE-Thesis" in next_steps
        assert "https://gitee.com/cwf818/ZUFE-Thesis" in next_steps
        assert "模板压缩包" in next_steps
        assert "已解压的完整模板目录" in next_steps


def test_prepare_workspace_requires_consent_then_archives_old_outputs():
    prepare_workspace = load_module("prepare_workspace")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        source = root / "incoming.docx"
        source.write_bytes(b"word source")
        target = root / "workspace/input/thesis.docx"

        result = prepare_workspace.prepare(
            root,
            source,
            move_word=False,
            copy_word=False,
            archive_existing=False,
        )
        assert result["status"] == "needs_confirmation"
        assert source.exists()
        assert not target.exists()

        result = prepare_workspace.prepare(
            root,
            source,
            move_word=False,
            copy_word=True,
            archive_existing=False,
        )
        assert result["status"] == "passed"
        assert source.exists()
        assert target.read_bytes() == b"word source"

        old_result = root / "workspace/output/qa_result.json"
        old_result.write_text("{}", encoding="utf-8")
        result = prepare_workspace.prepare(
            root,
            target,
            move_word=False,
            copy_word=False,
            archive_existing=False,
        )
        assert result["status"] == "blocked"
        assert old_result.exists()

        result = prepare_workspace.prepare(
            root,
            target,
            move_word=False,
            copy_word=False,
            archive_existing=True,
        )
        assert result["status"] == "passed"
        assert not old_result.exists()
        old_outputs = next(check for check in result["checks"] if check["name"] == "old_outputs")
        assert old_outputs["archived"]


def test_latex_escape_ascii_double_quotes_and_single_scan():
    common = load_module("common")
    assert common.latex_escape('"产品差异化"') == "``产品差异化''"
    assert common.latex_escape('A&B "test"') == r"A\&B ``test''"
    assert common.latex_escape("“中文引号”") == "“中文引号”"
    assert common.latex_escape("student's") == "student's"
    assert common.latex_escape(r"\alpha {x}") == r"\textbackslash{}alpha \{x\}"


def test_metadata_parser_preserves_numeric_identifiers_with_leading_zeroes():
    common = load_module("common")
    with tempfile.TemporaryDirectory() as tmp:
        metadata_path = Path(tmp) / "metadata.yaml"
        metadata_path.write_text(
            "student_id: 00123\nclass_name: 0007\nreport_style: 1\n",
            encoding="utf-8",
        )
        metadata = common.load_metadata_yaml(metadata_path)

        assert metadata["student_id"] == "00123"
        assert metadata["class_name"] == "0007"
        assert metadata["report_style"] == "1"


def test_import_docx_preserves_image_anchor_order():
    import_docx = load_module("import_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        png = root / "anchor.png"
        write_tiny_png(png)
        document = Document()
        document.add_paragraph("图片前段落")
        document.add_picture(str(png))
        document.add_paragraph("图片后段落")
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        import_docx.extract(root, docx_path)
        thesis = json.loads(
            (root / "workspace/intermediate/thesis.json").read_text(encoding="utf-8")
        )
        blocks = thesis["source_blocks"]
        before = next(block for block in blocks if block.get("text") == "图片前段落")
        after = next(block for block in blocks if block.get("text") == "图片后段落")
        image = next(block for block in blocks if block.get("source_type") == "image")

        assert before["order"] < image["order"] < after["order"]
        assert image["status"] == "needs_confirmation"
        assert image["asset_status"] == "pending_export"
        assert image["target_slot"] is None
        assert image["evidence"]["docx_media_path"].startswith("word/media/")
        assert image["evidence"]["anchor_paragraph_id"] == "p0002"
        assert thesis["counts"]["paragraphs"] == 3
        assert thesis["counts"]["source_blocks_by_type"]["paragraph"] == 2
        assert thesis["counts"]["source_blocks_by_type"]["image"] == 1


def test_import_docx_preserves_repeated_uses_of_the_same_image_relationship():
    import_docx = load_module("import_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        png = root / "repeated.png"
        write_tiny_png(png)
        document = Document()
        paragraph = document.add_paragraph("重复图片")
        run = paragraph.add_run()
        run.add_picture(str(png))
        run.add_picture(str(png))
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        import_docx.extract(root, docx_path)
        thesis = json.loads(
            (root / "workspace/intermediate/thesis.json").read_text(encoding="utf-8")
        )
        images = [block for block in thesis["source_blocks"] if block.get("source_type") == "image"]

        assert len(images) == 2
        assert images[0]["evidence"]["docx_media_path"] == images[1]["evidence"]["docx_media_path"]
        assert [image["evidence"]["anchor_image_occurrence"] for image in images] == [1, 2]


def test_export_assets_does_not_mark_image_semantic_position_mapped():
    import_docx = load_module("import_docx")
    export_assets = load_module("export_assets")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        png = root / "anchor.png"
        write_tiny_png(png)
        document = Document()
        document.add_paragraph("图片前段落")
        document.add_picture(str(png))
        document.add_paragraph("图片后段落")
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        import_docx.extract(root, docx_path)
        thesis_path = root / "workspace/intermediate/thesis.json"
        export_assets.export_assets(root, docx_path, thesis_path)
        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        image = next(
            block for block in thesis["source_blocks"] if block.get("source_type") == "image"
        )

        assert image["status"] == "needs_confirmation"
        assert image["target_slot"] is None
        assert image["asset_status"] == "exported"
        assert image["asset_output"].startswith("Images/word_media/")
        assert image["render_result"]["kind"] == "asset_extracted"


def test_import_docx_reports_unsupported_features():
    import_docx = load_module("import_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        document = Document()
        document.add_paragraph("正文段落")
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        insertion = (
            "<w:p><w:hyperlink><w:r><w:t>链接文本</w:t></w:r></w:hyperlink></w:p>"
            "<w:p><w:r><m:oMath><m:r><m:t>x=1</m:t></m:r></m:oMath></w:r></w:p>"
            '<w:p><w:ins w:id="1" w:author="tester"><w:r><w:t>修订文本</w:t></w:r></w:ins></w:p>'
            "<w:p><w:r><w:pict><v:textbox><w:txbxContent><w:p><w:r><w:t>文本框</w:t></w:r></w:p></w:txbxContent></v:textbox></w:pict></w:r></w:p>"
            "<w:p><w:pPr><w:numPr/></w:pPr><w:r><w:t>自动编号</w:t></w:r></w:p>"
            '<w:p><w:fldSimple w:instr="REF target"><w:r><w:t>域结果</w:t></w:r></w:fldSimple></w:p>'
            '<w:p><w:r><w:drawing><w:blip r:link="rId999"/></w:drawing></w:r></w:p>'
            "<w:p><w:r><w:drawing><chart/></w:drawing></w:r></w:p>"
            "<w:p><w:r><w:object><OLEObject/></w:object></w:r></w:p>"
        )
        rewrite_docx_xml(
            docx_path,
            {"word/document.xml": document_xml.replace("<w:sectPr", insertion + "<w:sectPr", 1)},
            {
                "word/footnotes.xml": (
                    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:footnote w:id="1"><w:p><w:r><w:t>脚注</w:t></w:r></w:p></w:footnote>'
                    "</w:footnotes>"
                ),
                "word/comments.xml": (
                    '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:comment w:id="1"><w:p><w:r><w:t>批注</w:t></w:r></w:p></w:comment>'
                    "</w:comments>"
                ),
                "word/header1.xml": (
                    '<w:hdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    "<w:p><w:r><w:t>页眉</w:t></w:r></w:p></w:hdr>"
                ),
            },
        )

        import_docx.extract(root, docx_path)
        thesis = json.loads(
            (root / "workspace/intermediate/thesis.json").read_text(encoding="utf-8")
        )
        features = {feature["type"]: feature for feature in thesis["unsupported_features"]}

        assert features["hyperlink"]["count"] == 1
        assert features["equation_omml"]["count"] == 1
        assert features["tracked_changes"]["count"] == 1
        assert features["textbox"]["count"] == 1
        assert features["footnote_or_endnote"]["count"] == 1
        assert features["comment"]["count"] == 1
        assert features["header_footer"]["count"] == 1
        assert features["linked_image"]["count"] == 1
        assert features["chart_or_smartart"]["count"] == 1
        assert features["ole_object"]["count"] >= 1
        assert features["field_code"]["count"] >= 1
        assert features["automatic_numbering"]["count"] == 1
        for feature in features.values():
            assert feature["status"] == "needs_confirmation"
            assert feature["locations"]


def test_import_and_prescan_expand_block_content_controls():
    import_docx = load_module("import_docx")
    prescan_docx = load_module("prescan_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        document = Document()
        document.add_paragraph("普通正文")
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
        content_control = (
            '<w:sdt><w:sdtPr><w:tag w:val="audit"/></w:sdtPr><w:sdtContent>'
            "<w:p><w:r><w:t>内容控件正文</w:t></w:r></w:p>"
            "</w:sdtContent></w:sdt>"
        )
        rewrite_docx_xml(
            docx_path,
            {
                "word/document.xml": document_xml.replace(
                    "<w:sectPr", content_control + "<w:sectPr", 1
                )
            },
        )

        prescan_result = prescan_docx.prescan(root, docx_path)
        assert prescan_result["status"] == "passed"
        assert prescan_result["counts"]["non_empty_paragraphs"] == 2
        assert any(block["text"] == "内容控件正文" for block in prescan_result["structure_preview"])

        import_docx.extract(root, docx_path)
        thesis = json.loads(
            (root / "workspace/intermediate/thesis.json").read_text(encoding="utf-8")
        )
        controlled = next(
            block for block in thesis["source_blocks"] if block.get("text") == "内容控件正文"
        )
        features = {feature["type"]: feature for feature in thesis["unsupported_features"]}
        assert controlled["evidence"]["inside_content_control"] is True
        assert features["content_control"]["count"] == 1


def test_export_assets_blocks_changed_source_docx():
    import_docx = load_module("import_docx")
    export_assets = load_module("export_assets")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        document = Document()
        document.add_paragraph("正文")
        docx_path = root / "workspace/input/thesis.docx"
        document.save(docx_path)

        import_docx.extract(root, docx_path)
        thesis_path = root / "workspace/intermediate/thesis.json"
        original_thesis = thesis_path.read_text(encoding="utf-8")
        docx_path.write_bytes(docx_path.read_bytes() + b"changed-after-import")

        result = export_assets.export_assets(root, docx_path, thesis_path)
        assert result["status"] == "blocked"
        assert result["gate"] == "source_docx_changed"
        assert not (root / "Images/word_media").exists()
        assert thesis_path.read_text(encoding="utf-8") == original_thesis


def test_flow_b_gate_blocks_unconfirmed_unsupported_features():
    check_flow_b_gate = load_module("check_flow_b_gate")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/basicinfo.tex").write_text("基本信息\n", encoding="utf-8")
        (root / "chapters/mainbody.tex").write_text("\\input{chapters/1_intro}\n", encoding="utf-8")
        (root / "chapters/1_intro.tex").write_text("正文\n", encoding="utf-8")
        (root / "Reference.bib").write_text("% empty\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "counts": {"total_source_blocks": 0},
                    "source_blocks": [],
                    "unsupported_features": [
                        {
                            "type": "equation_omml",
                            "count": 1,
                            "severity": "high",
                            "status": "needs_confirmation",
                            "locations": [{"part": "word/document.xml", "count": 1}],
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = check_flow_b_gate.check(root, thesis_path)
        assert result["status"] == "blocked"
        assert any(
            issue["check"] == "unsupported_feature_confirmation" for issue in result["issues"]
        )


def test_flow_b_gate_blocks_manual_figure_reference_numbers():
    check_flow_b_gate = load_module("check_flow_b_gate")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/basicinfo.tex").write_text("基本信息\n", encoding="utf-8")
        (root / "chapters/mainbody.tex").write_text("\\input{chapters/1_intro}\n", encoding="utf-8")
        (root / "chapters/1_intro.tex").write_text(
            "正文具体可见图2.1。\n"
            "\\begin{figure}[htbp]\n"
            "\\caption{年龄分布图}\n"
            "\\label{fig:age-distribution}\n"
            "\\end{figure}\n",
            encoding="utf-8",
        )
        (root / "Reference.bib").write_text("% empty\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "counts": {"total_source_blocks": 1},
                    "unsupported_features": [],
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "rendered",
                            "text": "正文具体可见图2.1。",
                            "target_slot": "chapters/1_intro.tex",
                            "render_result": {
                                "path": "chapters/1_intro.tex",
                                "kind": "chapter_tex",
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = check_flow_b_gate.check(root, thesis_path)
        assert result["status"] == "blocked"
        issue = next(
            issue
            for issue in result["issues"]
            if issue["check"] == "manual_cross_reference_numbers"
        )
        assert "图2.1" in issue["examples"][0]


def test_flow_b_gate_blocks_broken_latex_label_refs():
    check_flow_b_gate = load_module("check_flow_b_gate")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/basicinfo.tex").write_text("基本信息\n", encoding="utf-8")
        (root / "chapters/mainbody.tex").write_text("\\input{chapters/1_intro}\n", encoding="utf-8")
        (root / "chapters/1_intro.tex").write_text(
            "正文具体可见图~\\ref{fig:missing}。\n"
            "\\begin{figure}[htbp]\n"
            "\\caption{年龄分布图}\n"
            "\\label{fig:duplicate}\n"
            "\\end{figure}\n"
            "\\begin{figure}[htbp]\n"
            "\\caption{年龄分布图副本}\n"
            "\\label{fig:duplicate}\n"
            "\\end{figure}\n",
            encoding="utf-8",
        )
        (root / "Reference.bib").write_text("% empty\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "counts": {"total_source_blocks": 1},
                    "unsupported_features": [],
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "rendered",
                            "text": "正文具体可见图2.1。",
                            "target_slot": "chapters/1_intro.tex",
                            "render_result": {
                                "path": "chapters/1_intro.tex",
                                "kind": "chapter_tex",
                            },
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = check_flow_b_gate.check(root, thesis_path)
        assert result["status"] == "blocked"
        issues = {issue["check"]: issue for issue in result["issues"]}
        assert issues["duplicate_latex_labels"]["labels"] == ["fig:duplicate"]
        assert issues["undefined_latex_refs"]["labels"] == ["fig:missing"]


def test_check_env_reports_missing_required_latex_packages():
    check_env = load_module("check_env")
    original_command_exists = check_env.command_exists
    original_kpsewhich_exists = getattr(check_env, "kpsewhich_exists", None)
    try:
        check_env.command_exists = lambda _name: True
        check_env.kpsewhich_exists = lambda filename: filename != "gb7714-2015.bbx"
        result = check_env.check("latex")
    finally:
        check_env.command_exists = original_command_exists
        if original_kpsewhich_exists is not None:
            check_env.kpsewhich_exists = original_kpsewhich_exists

    checks = {check["name"]: check for check in result["checks"]}
    assert checks["tex_package_ctexbook.cls"]["status"] == "passed"
    assert checks["tex_package_biblatex.sty"]["status"] == "passed"
    assert checks["tex_package_gb7714-2015.bbx"]["status"] == "blocked"
    assert result["status"] == "blocked"
    issues = {issue["code"]: issue for issue in result["issues"]}
    assert issues["tex_core_file_missing"]["severity"] == "blocking"
    assert issues["tex_core_file_missing"]["repair_policy"] == "ask_user_before_install"
    assert issues["tex_core_file_missing"]["verify_command"].endswith("--stage latex")


def test_check_env_python_docx_hint_uses_short_timeout_and_mirror_fallback():
    check_env = load_module("check_env")
    original_find_spec = check_env.importlib.util.find_spec
    try:
        check_env.importlib.util.find_spec = lambda name: (
            None if name == "docx" else original_find_spec(name)
        )
        result = check_env.check("minimal")
    finally:
        check_env.importlib.util.find_spec = original_find_spec

    checks = {check["name"]: check for check in result["checks"]}
    hint = checks["python-docx"]["install_hint"]
    assert "--timeout 8" in hint
    assert "pypi.tuna.tsinghua.edu.cn/simple" in hint
    assert "失败、超时或无响应" in hint
    issues = {issue["code"]: issue for issue in result["issues"]}
    assert issues["python_docx_missing"]["severity"] == "blocking"
    assert issues["python_docx_missing"]["layer"] == "python-package"
    assert issues["python_docx_missing"]["verify_command"].endswith("--stage minimal")


def test_check_env_reports_missing_latex_commands_as_structured_issues():
    check_env = load_module("check_env")
    original_command_exists = check_env.command_exists
    original_kpsewhich_exists = getattr(check_env, "kpsewhich_exists", None)
    try:
        check_env.command_exists = lambda name: name not in {"xelatex", "biber"}
        check_env.kpsewhich_exists = lambda _filename: True
        result = check_env.check("latex")
    finally:
        check_env.command_exists = original_command_exists
        if original_kpsewhich_exists is not None:
            check_env.kpsewhich_exists = original_kpsewhich_exists

    issues = result["issues"]
    assert result["status"] == "blocked"
    assert [issue["code"] for issue in issues] == [
        "tex_command_missing",
        "tex_command_missing",
    ]
    assert {issue["target"] for issue in issues} == {"xelatex", "biber"}
    assert all(issue["verify_command"].endswith("--stage latex") for issue in issues)


def test_check_env_qa_stage_reports_optional_tools_without_blocking():
    check_env = load_module("check_env")
    original_command_exists = check_env.command_exists
    try:
        check_env.command_exists = lambda _name: False
        result = check_env.check("qa")
    finally:
        check_env.command_exists = original_command_exists

    checks = {check["name"]: check for check in result["checks"]}
    assert checks["pdfinfo"]["status"] == "needs_review"
    assert checks["pdftotext"]["status"] == "needs_review"
    assert result["status"] == "needs_review"
    assert {issue["code"] for issue in result["issues"]} == {"qa_tool_missing"}
    assert all(issue["severity"] == "optional" for issue in result["issues"])


def test_check_env_main_allows_needs_review_exit_code():
    check_env = load_module("check_env")
    original_check = check_env.check
    try:
        check_env.check = lambda _stage: {
            "status": "needs_review",
            "checks": [],
            "issues": [],
        }
        with contextlib.redirect_stdout(io.StringIO()):
            assert check_env.main([]) == 0
    finally:
        check_env.check = original_check


def test_check_env_distinguishes_missing_kpsewhich_from_tex_packages():
    check_env = load_module("check_env")
    with patch.object(check_env, "command_exists", side_effect=lambda name: name != "kpsewhich"):
        result = check_env.check("latex")

    checks = {check["name"]: check for check in result["checks"]}
    assert checks["kpsewhich"]["status"] == "blocked"
    assert not any(name.startswith("tex_package_") for name in checks)
    assert any(
        issue["code"] == "tex_command_missing" and issue["target"] == "kpsewhich"
        for issue in result["issues"]
    )
    assert result["scope"]["not_checked"]


def test_check_env_blocks_unsupported_python_version():
    check_env = load_module("check_env")
    with (
        patch.object(check_env.sys, "version_info", (3, 9, 18)),
        patch.object(check_env, "command_exists", return_value=True),
    ):
        result = check_env.check("qa")

    assert result["status"] == "blocked"
    assert result["checks"][0]["status"] == "blocked"
    assert any(issue["code"] == "python_version_unsupported" for issue in result["issues"])


def test_check_env_detects_broken_python_docx_import():
    check_env = load_module("check_env")
    with (
        patch.object(check_env.importlib.util, "find_spec", return_value=object()),
        patch.object(
            check_env.importlib,
            "import_module",
            side_effect=ImportError("broken lxml dependency"),
        ),
        patch.object(check_env, "pip_available", return_value=True),
    ):
        result = check_env.check("minimal")

    assert result["status"] == "blocked"
    issue_codes = {issue["code"] for issue in result["issues"]}
    assert "python_docx_import_failed" in issue_codes
    assert "python_docx_missing" not in issue_codes


def test_check_env_rejects_unrelated_docx_module_without_document_api():
    check_env = load_module("check_env")
    with (
        patch.object(check_env.importlib.util, "find_spec", return_value=object()),
        patch.object(check_env.importlib, "import_module", return_value=object()),
        patch.object(check_env, "pip_available", return_value=True),
    ):
        result = check_env.check("minimal")

    python_docx_check = next(check for check in result["checks"] if check["name"] == "python-docx")
    assert result["status"] == "blocked"
    assert "Document API" in python_docx_check["detail"]
    assert {issue["code"] for issue in result["issues"]} >= {"python_docx_import_failed"}


def test_prescan_reads_cover_table_metadata_without_report_style_default():
    prescan_docx = load_module("prescan_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        docx_path = root / "cover.docx"
        document = Document()
        document.add_paragraph("专业实践报告")
        table = document.add_table(rows=4, cols=2)
        rows = [
            ("指导教师", "张老师"),
            ("专业名称", "数字经济"),
            ("学院", "经济学院"),
            ("日期", "2026年6月"),
        ]
        for row, (label, value) in zip(table.rows, rows, strict=True):
            row.cells[0].text = label
            row.cells[1].text = value
        document.save(docx_path)

        result = prescan_docx.prescan(root, docx_path)
        candidates = result["metadata_candidates"]
        assert candidates["report_style"] == "1"
        assert candidates["mentor"] == "张老师"
        assert candidates["major"] == "数字经济"
        assert candidates["college"] == "经济学院"
        assert candidates["date"] == "2026年6月"

        assert prescan_docx.metadata_candidates(["普通论文标题"])["report_style"] == ""


def test_prescan_and_import_accept_table_only_docx_metadata():
    import_docx = load_module("import_docx")
    prescan_docx = load_module("prescan_docx")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/input").mkdir(parents=True)
        docx_path = root / "workspace/input/thesis.docx"
        document = Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "学生姓名"
        table.cell(0, 1).text = "张三"
        table.cell(1, 0).text = "学号"
        table.cell(1, 1).text = "20260001"
        document.save(docx_path)

        prescan_result = prescan_docx.prescan(root, docx_path)
        import_result = import_docx.extract(root, docx_path)
        thesis = json.loads(
            (root / "workspace/intermediate/thesis.json").read_text(encoding="utf-8")
        )

        assert prescan_result["status"] == "passed"
        assert prescan_result["counts"]["non_empty_paragraphs"] == 0
        assert prescan_result["metadata_candidates"]["name"] == "张三"
        assert import_result["counts"]["tables"] == 1
        assert thesis["metadata_candidates"]["name"] == "张三"
        assert thesis["metadata_candidates"]["student_id"] == "20260001"


def test_render_basicinfo_blocks_missing_report_style():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        metadata = root / "metadata.yaml"
        metadata.write_text("thesis_title_cn: 测试题目\n", encoding="utf-8")

        result = render_basicinfo.render(root, metadata, thesis_path=None)
        assert result["status"] == "blocked"
        assert "report_style" in result["missing_fields"]
        assert not (root / "chapters/basicinfo.tex").exists()


def test_render_basicinfo_blocks_missing_required_cover_metadata():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        metadata = root / "metadata.yaml"
        metadata.write_text(
            "report_style: 1\nenglish_content_decision: omit\n",
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path=None)
        assert result["status"] == "blocked"
        assert result["gate"] == "metadata_required"
        assert "thesis_title_cn" in result["missing_fields"]
        assert "name" in result["missing_fields"]
        assert "student_id" in result["missing_fields"]
        assert not (root / "chapters/basicinfo.tex").exists()


def test_render_basicinfo_blocks_missing_subtitle_when_enabled():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(
                extra="has_subtitle: true\nenglish_content_decision: omit\n",
            ),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path=None)
        assert result["status"] == "blocked"
        assert result["gate"] == "metadata_required"
        assert "thesis_subtitle_cn" in result["missing_fields"]
        assert "thesis_subtitle_en" in result["missing_fields"]
        assert not (root / "chapters/basicinfo.tex").exists()


def test_render_basicinfo_blocks_unapproved_generated_english():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(basic_metadata_yaml(), encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "metadata": {
                        "abstract_en": "Generated English abstract.",
                        "keywords_en": ["generated", "keywords"],
                        "english_content_source": "generated",
                    }
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "blocked"
        assert result["gate"] == "generated_english_requires_confirmation"
        assert not (root / "chapters/basicinfo.tex").exists()

        metadata.write_text(
            basic_metadata_yaml(extra="allow_generated_english: true\n"),
            encoding="utf-8",
        )
        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "passed"


def test_render_basicinfo_requires_missing_english_content_decision():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(basic_metadata_yaml(), encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps({"metadata": {"abstract_cn": "中文摘要。"}}, ensure_ascii=False),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "blocked"
        assert result["gate"] == "english_content_decision_required"
        assert not (root / "chapters/basicinfo.tex").exists()

        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "passed"


def test_render_basicinfo_requires_explicit_source_field_evidence():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "metadata": {},
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "mapped",
                            "source_type": "paragraph",
                            "text": "绝不能丢失的封面附注",
                            "target_slot": "chapters/basicinfo.tex",
                            "render_result": None,
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        assert result["status"] == "needs_confirmation"
        assert result["gate"] == "basicinfo_source_evidence"
        assert result["unverified_blocks"][0]["reasons"] == ["missing_metadata_fields"]
        assert thesis["source_blocks"][0]["status"] == "mapped"
        assert thesis["source_blocks"][0]["render_result"] is None
        assert "绝不能丢失的封面附注" not in (root / "chapters/basicinfo.tex").read_text(
            encoding="utf-8"
        )


def test_render_basicinfo_marks_only_verified_metadata_bindings_rendered():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "metadata": {},
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "mapped",
                            "source_type": "paragraph",
                            "text": "学生姓名：测试姓名",
                            "metadata_fields": ["name"],
                            "target_slot": "chapters/basicinfo.tex",
                            "render_result": None,
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        block = thesis["source_blocks"][0]
        assert result["status"] == "passed"
        assert block["status"] == "rendered"
        assert block["render_result"]["metadata_fields"] == ["name"]
        assert block["render_result"]["evidence"] == "all_bound_values_found_in_source_block"


def test_render_basicinfo_requires_an_explicit_destination_for_residual_text():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "metadata": {},
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "mapped",
                            "source_type": "paragraph",
                            "text": "学生姓名：测试姓名；附注：不得公开",
                            "metadata_fields": ["name"],
                            "target_slot": "chapters/basicinfo.tex",
                            "render_result": None,
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "needs_confirmation"
        assert any(
            reason.startswith("uncovered_source_text:")
            for reason in result["unverified_blocks"][0]["reasons"]
        )

        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        block = thesis["source_blocks"][0]
        block["metadata_excluded_text"] = ["附注：不得公开"]
        block["metadata_exclusion_reason"] = "隐私提示，不属于论文封面输出。"
        thesis_path.write_text(json.dumps(thesis, ensure_ascii=False), encoding="utf-8")

        result = render_basicinfo.render(root, metadata, thesis_path)
        rendered = json.loads(thesis_path.read_text(encoding="utf-8"))["source_blocks"][0]
        assert result["status"] == "passed"
        assert rendered["status"] == "rendered"


def test_render_basicinfo_accepts_report_style_candidate_synonyms():
    render_basicinfo = load_module("render_basicinfo")
    verified, problems = render_basicinfo.verify_basicinfo_block(
        {
            "text": "实践报告",
            "metadata_fields": ["report_style"],
        },
        {"report_style": "1"},
    )
    assert verified == ["report_style"]
    assert problems == []


def test_render_basicinfo_blocks_malformed_ledger_without_writing():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps({"metadata": {}, "source_blocks": {}}, ensure_ascii=False),
            encoding="utf-8",
        )

        result = render_basicinfo.render(root, metadata, thesis_path)
        assert result["status"] == "blocked"
        assert result["gate"] == "thesis_json_structure"
        assert not (root / "chapters/basicinfo.tex").exists()


def test_qa_flags_bibtex_and_citation_lint_failures():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/1_intro.tex").write_text(
            r"正文引用 \cite{known,missing}。",
            encoding="utf-8",
        )
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/1_intro}\n",
            encoding="utf-8",
        )
        (root / "Reference.bib").write_text(
            "@article{known,\n"
            "  title={A}\n"
            "}\n"
            "@book{known,\n"
            "  title={B}\n"
            "}\n"
            "@misc{broken,\n"
            "  title={Broken}\n",
            encoding="utf-8",
        )
        (root / "workspace/intermediate/thesis.json").write_text("{}", encoding="utf-8")

        checks = {check["name"]: check for check in qa.source_quality_checks(root)}
        assert checks["bibtex_duplicate_keys"]["status"] == "failed"
        assert "known" in checks["bibtex_duplicate_keys"]["detail"]
        assert checks["bibtex_braces_balanced"]["status"] == "failed"
        assert checks["citation_keys_defined"]["status"] == "failed"
        assert "missing" in checks["citation_keys_defined"]["detail"]


def test_qa_ignores_stale_chapters_not_referenced_by_current_mainbody():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/active}\n",
            encoding="utf-8",
        )
        (root / "chapters/active.tex").write_text("当前正文。\n", encoding="utf-8")
        (root / "chapters/stale.tex").write_text(
            "xxxxxxxxxxxx\\label{duplicate}\\label{duplicate}\\ref{missing}\n",
            encoding="utf-8",
        )
        thesis = {"structure": {"chapters": []}, "source_blocks": []}
        (root / "workspace/intermediate/thesis.json").write_text(
            json.dumps(thesis),
            encoding="utf-8",
        )

        source_text = qa.rendered_source_text(root, thesis)
        checks = {check["name"]: check for check in qa.source_quality_checks(root, thesis)}

        assert "当前正文" in source_text
        assert "xxxxxxxxxxxx" not in source_text
        assert checks["source_duplicate_latex_labels"]["status"] == "passed"
        assert checks["source_undefined_latex_refs"]["status"] == "passed"


def test_qa_reads_parenthesized_bibtex_keys():
    qa = load_module("qa")
    assert qa.bibtex_keys("@article(parenthesized, title={A})\n") == ["parenthesized"]


def test_render_bib_preserves_existing_file_until_mapping_is_confirmed():
    render_bib = load_module("render_bib")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/intermediate").mkdir(parents=True)
        target = root / "Reference.bib"
        target.write_text("% existing bibliography\n", encoding="utf-8")
        empty_bib = root / "workspace/input/references.bib"
        empty_bib.parent.mkdir(parents=True)
        empty_bib.write_text("\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "references": [],
                    "source_blocks": [
                        {
                            "id": "r0001",
                            "status": "needs_confirmation",
                            "target_slot": "Reference.bib",
                            "bibtex": "@article{confirmed, title={Confirmed}}",
                        }
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_bib.render(root, thesis_path, input_bib=empty_bib)
        assert result["status"] == "needs_confirmation"
        assert result["target_written"] is False
        assert any("外部 BibTeX 文件为空" in warning for warning in result["warnings"])
        assert target.read_text(encoding="utf-8") == "% existing bibliography\n"

        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        thesis["source_blocks"][0]["status"] = "mapped"
        thesis_path.write_text(json.dumps(thesis, ensure_ascii=False), encoding="utf-8")
        result = render_bib.render(root, thesis_path, input_bib=None)
        rendered = json.loads(thesis_path.read_text(encoding="utf-8"))

        assert result["status"] == "passed"
        assert result["target_written"] is True
        assert "@article{confirmed" in target.read_text(encoding="utf-8")
        assert rendered["source_blocks"][0]["status"] == "rendered"
        assert rendered["source_blocks"][0]["render_result"]["path"] == "Reference.bib"


def test_render_bib_does_not_partially_write_when_any_entry_is_unresolved():
    render_bib = load_module("render_bib")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/intermediate").mkdir(parents=True)
        target = root / "Reference.bib"
        target.write_text("% existing bibliography\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "references": [],
                    "source_blocks": [
                        {
                            "id": "r0001",
                            "status": "mapped",
                            "target_slot": "Reference.bib",
                            "bibtex": "@article{confirmed, title={Confirmed}}",
                        },
                        {
                            "id": "r0002",
                            "status": "mapped",
                            "target_slot": "Reference.bib",
                            "text": "尚未转换的参考文献",
                        },
                    ],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_bib.render(root, thesis_path, input_bib=None)
        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))
        assert result["status"] == "needs_confirmation"
        assert result["target_written"] is False
        assert target.read_text(encoding="utf-8") == "% existing bibliography\n"
        assert thesis["source_blocks"][0]["status"] == "mapped"
        assert thesis["source_blocks"][1]["status"] == "needs_confirmation"

        thesis["source_blocks"][1]["status"] = "discarded_with_reason"
        thesis["source_blocks"][1]["discard_reason"] = "用户确认该行不是参考文献"
        thesis_path.write_text(json.dumps(thesis, ensure_ascii=False), encoding="utf-8")
        result = render_bib.render(root, thesis_path, input_bib=None)

        assert result["status"] == "passed"
        assert result["target_written"] is True
        assert "@article{confirmed" in target.read_text(encoding="utf-8")


def test_render_chapters_blocks_prefix_path_escape():
    render_chapters = load_module("render_chapters")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "mapped",
                            "text": "越界正文",
                            "target_slot": "chapters_evil/escape.tex",
                        }
                    ],
                    "structure": {
                        "chapters": [
                            {
                                "title": "bad",
                                "file": "chapters_evil/escape.tex",
                                "block_ids": ["p0001"],
                            }
                        ]
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        result = render_chapters.render(root, thesis_path, allow_incomplete=False)
        assert result["status"] == "blocked"
        assert not (root / "chapters_evil/escape.tex").exists()


def test_render_chapters_blocks_duplicate_chapter_targets_before_writing():
    render_chapters = load_module("render_chapters")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "mapped",
                            "source_type": "paragraph",
                            "text": "第一段",
                            "target_slot": "chapters/1_intro.tex",
                        },
                        {
                            "id": "p0002",
                            "status": "mapped",
                            "source_type": "paragraph",
                            "text": "第二段",
                            "target_slot": "chapters/1_intro.tex",
                        },
                    ],
                    "structure": {
                        "chapters": [
                            {
                                "title": "第一章",
                                "file": "chapters/1_intro.tex",
                                "block_ids": ["p0001"],
                            },
                            {
                                "title": "重复目标",
                                "file": "chapters/./1_intro.tex",
                                "block_ids": ["p0002"],
                            },
                        ]
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_chapters.render(root, thesis_path, allow_incomplete=False)

        assert result["status"] == "blocked"
        assert "duplicate_chapter_file" in {issue["check"] for issue in result["issues"]}
        assert not (root / "chapters/1_intro.tex").exists()


def test_render_chapters_requires_an_exported_existing_image_asset():
    render_chapters = load_module("render_chapters")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis = {
            "source_blocks": [
                {
                    "id": "i0001",
                    "status": "mapped",
                    "source_type": "image",
                    "summary": "示例图片",
                    "target_slot": "chapters/1_intro.tex",
                }
            ],
            "structure": {
                "chapters": [
                    {
                        "title": "第一章",
                        "file": "chapters/1_intro.tex",
                        "block_ids": ["i0001"],
                    }
                ]
            },
        }
        thesis_path.write_text(json.dumps(thesis, ensure_ascii=False), encoding="utf-8")

        result = render_chapters.render(root, thesis_path, allow_incomplete=False)
        assert result["status"] == "blocked"
        assert result["invalid_structure_blocks"][0]["status"] == "image_asset_invalid"
        assert not (root / "chapters/1_intro.tex").exists()

        asset = root / "Images/word_media/image1.png"
        asset.parent.mkdir(parents=True)
        asset.write_bytes(TINY_PNG)
        thesis["source_blocks"][0]["asset_status"] = "exported"
        thesis["source_blocks"][0]["asset_output"] = "Images/word_media/image1.png"
        thesis_path.write_text(json.dumps(thesis, ensure_ascii=False), encoding="utf-8")

        result = render_chapters.render(root, thesis_path, allow_incomplete=False)
        chapter_text = (root / "chapters/1_intro.tex").read_text(encoding="utf-8")
        assert result["status"] == "passed"
        assert "Images/word\\_media/image1.png" in chapter_text
        assert "chapters/1_intro.tex}" not in chapter_text


def test_render_chapters_blocks_discarded_block_in_structure():
    render_chapters = load_module("render_chapters")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "source_blocks": [
                        {
                            "id": "p0001",
                            "status": "discarded_with_reason",
                            "discard_reason": "模板说明文字",
                            "text": "不应该被写入正文",
                            "target_slot": "chapters/1_intro.tex",
                        }
                    ],
                    "structure": {
                        "chapters": [
                            {
                                "title": "intro",
                                "file": "chapters/1_intro.tex",
                                "block_ids": ["p0001"],
                            }
                        ]
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        result = render_chapters.render(root, thesis_path, allow_incomplete=False)
        thesis = json.loads(thesis_path.read_text(encoding="utf-8"))

        assert result["status"] == "blocked"
        assert result["blocking_blocks"] == ["p0001"]
        assert not (root / "chapters/1_intro.tex").exists()
        assert thesis["source_blocks"][0]["status"] == "discarded_with_reason"
        assert thesis["source_blocks"][0].get("render_result") is None


def test_render_basicinfo_supports_thesis_title_abs():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(
                title_cn="封面题目",
                extra="thesis_title_abs_cn: 摘要页题目\nenglish_content_decision: omit\n",
            ),
            encoding="utf-8",
        )
        render_basicinfo.render(root, metadata, thesis_path=None)
        basicinfo = (root / "chapters/basicinfo.tex").read_text(encoding="utf-8")
        assert "\\newcommand{\\thesisTitle}{封面题目}" in basicinfo
        assert "\\newcommand{\\thesisTitleAbs}{摘要页题目}" in basicinfo


def test_render_basicinfo_hides_hyperref_link_borders():
    render_basicinfo = load_module("render_basicinfo")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        metadata = root / "metadata.yaml"
        metadata.write_text(
            basic_metadata_yaml(extra="english_content_decision: omit\n"),
            encoding="utf-8",
        )
        render_basicinfo.render(root, metadata, thesis_path=None)
        basicinfo = (root / "chapters/basicinfo.tex").read_text(encoding="utf-8")
        assert r"\hypersetup{hidelinks,pdfborder={0 0 0},pdfborderstyle={/S/U/W 0}}" in basicinfo


def test_build_xelatex_uses_noninteractive_error_flags():
    build = load_module("build")
    xelatex_steps = [command for command in build.COMPILE_CHAIN if command[0] == "xelatex"]
    assert xelatex_steps
    for command in xelatex_steps:
        assert command[:4] == [
            "xelatex",
            "-interaction=nonstopmode",
            "-halt-on-error",
            "-file-line-error",
        ]
        assert command[-1] == "main.tex"


def test_build_timeout_preserves_byte_output_without_type_error():
    build = load_module("build")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        timeout = subprocess.TimeoutExpired(
            cmd=build.COMPILE_CHAIN[0],
            timeout=1,
            output=b"partial build output\n",
        )
        with patch.object(build.subprocess, "run", side_effect=timeout):
            steps = build.run_chain(root, timeout=1)

        assert steps[0]["exit_code"] == 124
        log = (root / steps[0]["log"]).read_text(encoding="utf-8")
        assert "partial build output" in log
        assert "command timed out after 1 seconds" in log


def test_build_success_writes_machine_and_human_reports():
    build = load_module("build")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/output").mkdir(parents=True)

        def successful_chain(project_root: Path, timeout: int) -> list[dict]:
            assert timeout == 30
            (project_root / "main.pdf").write_bytes(b"%PDF-1.7\n")
            return [
                {
                    "index": index,
                    "command": command,
                    "exit_code": 0,
                    "started_at": "2026-01-01T00:00:00Z",
                    "ended_at": "2026-01-01T00:00:01Z",
                    "log": f"workspace/output/build-step-{index}.log",
                }
                for index, command in enumerate(build.COMPILE_CHAIN, start=1)
            ]

        with (
            patch.object(
                build,
                "check_flow_b_gate",
                return_value={
                    "status": "passed",
                    "thesis_json_fingerprint": {"sha256": "thesis", "size_bytes": 1},
                    "source_docx_fingerprint": {"sha256": "docx", "size_bytes": 1},
                },
            ),
            patch.object(build, "prepare_build", return_value=[]),
            patch.object(build, "run_chain", side_effect=successful_chain),
        ):
            result = build.build(root, timeout=30)

        assert result["status"] == "passed"
        assert result["new_pdf"] is True
        machine_report = json.loads(
            (root / "workspace/output/build_result.json").read_text(encoding="utf-8")
        )
        human_report = (root / "workspace/output/report.md").read_text(encoding="utf-8")
        assert machine_report["status"] == "passed"
        assert len(machine_report["steps"]) == len(build.COMPILE_CHAIN)
        assert "- Status: `passed`" in human_report


def test_build_stops_before_archiving_when_flow_b_gate_is_blocked():
    build = load_module("build")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        old_pdf = root / "main.pdf"
        old_pdf.write_bytes(b"old pdf")
        with (
            patch.object(
                build,
                "check_flow_b_gate",
                return_value={
                    "status": "blocked",
                    "issues": [{"check": "source_docx_changed"}],
                },
            ),
            patch.object(build, "prepare_build") as prepare_build,
            patch.object(build, "run_chain") as run_chain,
        ):
            result = build.build(root, timeout=30)

        assert result["status"] == "blocked"
        assert result["steps"] == []
        assert old_pdf.read_bytes() == b"old pdf"
        prepare_build.assert_not_called()
        run_chain.assert_not_called()
        assert (root / "workspace/output/build_result.json").exists()


def test_diagnose_build_classifies_standard_latex_missing_file_messages():
    diagnose_build = load_module("diagnose_build")
    issues = diagnose_build.classify(
        "! LaTeX Error: File `missing-package.sty' not found.\n"
        "! LaTeX Error: File `Images/missing-figure.png' not found.\n"
    )
    categories = {issue["category"] for issue in issues}
    assert "environment_issue" in categories
    assert "user_input_required" in categories


def test_diagnose_build_writes_actionable_result():
    diagnose_build = load_module("diagnose_build")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/output").mkdir(parents=True)
        (root / "main.log").write_text(
            "! LaTeX Error: File `missing-package.sty' not found.\n",
            encoding="utf-8",
        )

        result = diagnose_build.diagnose(root)
        written = json.loads((root / "workspace/output/diagnosis.json").read_text(encoding="utf-8"))

        assert result["status"] == "needs_action"
        assert result["logs_checked"] == ["main.log"]
        assert result["issues"][0]["category"] == "environment_issue"
        assert written == result


def test_flow_b_gate_validates_ledger_identity_and_chapter_ownership():
    check_flow_b_gate = load_module("check_flow_b_gate")
    thesis = {
        "counts": {"total_source_blocks": 2, "paragraphs": 2},
        "source_blocks": [
            {
                "id": "p0001",
                "source_type": "paragraph",
                "status": "mapped",
                "target_slot": "chapters/1_intro.tex",
            },
            {
                "id": "p0001",
                "source_type": "paragraph",
                "status": "unknown",
                "target_slot": "chapters/2_method.tex",
            },
        ],
        "structure": {
            "chapters": [
                {
                    "title": "保留文件误用",
                    "file": "chapters/mainbody.tex",
                    "block_ids": ["p0001", "missing"],
                },
                {
                    "title": "重复归属",
                    "file": "chapters/2_method.tex",
                    "block_ids": ["p0001"],
                },
            ]
        },
    }

    issue_names = {issue["check"] for issue in check_flow_b_gate.ledger_schema_issues(thesis)}
    assert "duplicate_source_block_id" in issue_names
    assert "source_block_status_value" in issue_names
    assert "chapter_file" in issue_names
    assert "chapter_unknown_block" in issue_names
    assert "chapter_duplicate_block_reference" in issue_names
    assert "chapter_target_mismatch" in issue_names


def test_flow_b_gate_reports_malformed_source_block_without_crashing():
    check_flow_b_gate = load_module("check_flow_b_gate")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/intermediate").mkdir(parents=True)
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "counts": {"total_source_blocks": 2},
                    "source_blocks": [
                        None,
                        {
                            "id": "p0001",
                            "status": [],
                            "source_type": {},
                            "text": "畸形状态源块",
                        },
                    ],
                    "structure": {"chapters": []},
                    "unsupported_features": [{"type": "field_code", "count": 1, "status": []}],
                }
            ),
            encoding="utf-8",
        )

        result = check_flow_b_gate.check(root, thesis_path)

        assert result["status"] == "blocked"
        issue_names = {issue["check"] for issue in result["issues"]}
        assert "source_block_type" in issue_names
        assert "source_block_status_value" in issue_names
        assert "unsupported_feature_confirmation" in issue_names


def test_flow_b_gate_rechecks_source_docx_fingerprint_before_flow_c():
    check_flow_b_gate = load_module("check_flow_b_gate")
    common = load_module("common")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/input").mkdir(parents=True)
        (root / "workspace/intermediate").mkdir(parents=True)
        source = root / "workspace/input/thesis.docx"
        source.write_bytes(b"source version one")
        (root / "chapters/basicinfo.tex").write_text("基本信息\n", encoding="utf-8")
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/1_intro}\n",
            encoding="utf-8",
        )
        (root / "chapters/1_intro.tex").write_text("正文\n", encoding="utf-8")
        (root / "Reference.bib").write_text("% empty\n", encoding="utf-8")
        thesis_path = root / "workspace/intermediate/thesis.json"
        thesis_path.write_text(
            json.dumps(
                {
                    "source_docx": "workspace/input/thesis.docx",
                    "source_docx_fingerprint": common.file_fingerprint(source),
                    "counts": {"total_source_blocks": 0},
                    "source_blocks": [],
                    "structure": {"chapters": []},
                    "unsupported_features": [],
                }
            ),
            encoding="utf-8",
        )

        assert check_flow_b_gate.check(root, thesis_path)["status"] == "passed"

        source.write_bytes(b"source version two")
        result = check_flow_b_gate.check(root, thesis_path)
        assert result["status"] == "blocked"
        assert "source_docx_changed" in {issue["check"] for issue in result["issues"]}


def test_render_chapters_table_uses_fixed_font_without_resizebox():
    render_chapters = load_module("render_chapters")
    latex = render_chapters.block_to_latex(
        {
            "source_type": "table",
            "table": {"rows": [["指标", "值"], ["样本", "1"]]},
        }
    )
    assert "\\zihao{5}" in latex
    assert "\\songti" in latex
    assert "\\resizebox" not in latex
    assert "\\begin{tabular}{@{}ll@{}}" in latex


def test_qa_flags_missing_superscript_rendering_and_resizebox():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/1_test.tex").write_text(
            "\\resizebox{\\textwidth}{!}{bad table}\n引用1\n具体可见图2.1。\n",
            encoding="utf-8",
        )
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/1_test}\n",
            encoding="utf-8",
        )
        thesis = {
            "source_blocks": [
                {
                    "id": "p0001",
                    "status": "rendered",
                    "runs": [
                        {"text": "引用", "superscript": False},
                        {"text": "1", "superscript": True},
                    ],
                }
            ]
        }
        (root / "workspace/intermediate/thesis.json").write_text(
            json.dumps(thesis, ensure_ascii=False),
            encoding="utf-8",
        )
        checks = {check["name"]: check for check in qa.source_quality_checks(root)}
        assert checks["source_table_resizebox_textwidth"]["status"] == "warning"
        assert checks["source_manual_cross_reference_numbers"]["status"] == "failed"
        assert "图2.1" in checks["source_manual_cross_reference_numbers"]["detail"]
        assert checks["source_duplicate_latex_labels"]["status"] == "passed"
        assert checks["source_undefined_latex_refs"]["status"] == "passed"
        assert checks["source_superscript_runs_rendered"]["status"] == "warning"


def test_qa_flags_broken_latex_label_refs():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/1_test.tex").write_text(
            "正文见图~\\ref{fig:missing}。\n"
            "\\caption{图一}\\label{fig:duplicate}\n"
            "\\caption{图二}\\label{fig:duplicate}\n",
            encoding="utf-8",
        )
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/1_test}\n",
            encoding="utf-8",
        )
        (root / "workspace/intermediate/thesis.json").write_text("{}", encoding="utf-8")

        checks = {check["name"]: check for check in qa.source_quality_checks(root)}
        assert checks["source_duplicate_latex_labels"]["status"] == "failed"
        assert checks["source_duplicate_latex_labels"]["detail"] == "fig:duplicate"
        assert checks["source_undefined_latex_refs"]["status"] == "failed"
        assert checks["source_undefined_latex_refs"]["detail"] == "fig:missing"


def test_qa_placeholder_scan_includes_generated_chapter_files():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "chapters/basicinfo.tex").write_text("基本信息\n", encoding="utf-8")
        (root / "chapters/mainbody.tex").write_text(
            "\\input{chapters/1_intro}\n",
            encoding="utf-8",
        )
        (root / "chapters/1_intro.tex").write_text(
            "正文里残留 xxxxxxxxxxxx\n",
            encoding="utf-8",
        )
        (root / "workspace/intermediate/thesis.json").write_text("{}", encoding="utf-8")
        result = qa.qa(root)
        checks = {check["name"]: check for check in result["checks"]}
        assert checks[r"placeholder_xxxxxxxxxxxx"]["status"] == "warning"


def test_qa_requires_build_result_for_pdf_freshness():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "main.pdf").write_bytes(b"%PDF-1.7\n1 0 obj << /Type /Page >> endobj\n")
        (root / "workspace/intermediate/thesis.json").write_text(
            "{}",
            encoding="utf-8",
        )

        result = qa.qa(root)
        checks = {check["name"]: check for check in result["checks"]}
        assert checks["pdf_exists"]["status"] == "passed"
        assert checks["pdf_freshness"]["status"] == "failed"
        assert "build_result.json" in checks["pdf_freshness"]["detail"]
        assert result["status"] == "failed"


def test_qa_reports_malformed_json_inputs_without_crashing():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "workspace/output").mkdir(parents=True)
        (root / "workspace/intermediate/thesis.json").write_text("{broken", encoding="utf-8")
        (root / "workspace/output/build_result.json").write_text("[broken", encoding="utf-8")

        result = qa.qa(root)
        checks = {check["name"]: check for check in result["checks"]}

        assert result["status"] == "failed"
        assert checks["flow_b_gate_current"]["status"] == "failed"
        assert checks["build_chain_passed"]["status"] == "failed"


def test_qa_unverifiable_page_count_requires_review_instead_of_false_failure():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "main.pdf").write_bytes(b"%PDF-1.7\ncompressed objects\n")
        (root / "workspace/intermediate/thesis.json").write_text("{}", encoding="utf-8")

        with patch.object(qa, "count_pdf_pages", return_value=0):
            result = qa.qa(root)

        page_check = next(check for check in result["checks"] if check["name"] == "page_count")
        assert page_check["status"] == "warning"
        assert "无法可靠确认" in page_check["detail"]


def test_qa_counts_pages_with_pdfinfo_before_byte_fallback():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        pdf = Path(tmp) / "main.pdf"
        pdf.write_bytes(b"%PDF-1.7\ncompressed page objects without plain markers\n")

        completed = subprocess.CompletedProcess(
            args=["pdfinfo", str(pdf)],
            returncode=0,
            stdout="Title: test\nPages:          15\n",
        )
        with (
            patch.object(qa.shutil, "which", return_value="/usr/bin/pdfinfo"),
            patch.object(qa.subprocess, "run", return_value=completed),
        ):
            assert qa.count_pdf_pages(pdf) == 15


def test_qa_tool_timeouts_degrade_without_crashing():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        pdf = Path(tmp) / "main.pdf"
        pdf.write_bytes(b"%PDF-1.7\n")
        timeout = subprocess.TimeoutExpired(cmd=["pdf-tool"], timeout=30)

        with (
            patch.object(qa.shutil, "which", return_value="/usr/bin/pdf-tool"),
            patch.object(qa.subprocess, "run", side_effect=timeout),
        ):
            assert qa.extract_text_with_pdftotext(pdf) == ""
            assert qa.count_pdf_pages_with_pdfinfo(pdf) == 0


def test_qa_body_signal_tolerates_malformed_ledger_sections():
    qa = load_module("qa")
    found, detail = qa.body_signal(
        {"structure": [], "source_blocks": {}},
        "正文文本",
    )
    assert found is False
    assert "no audited" in detail


def test_qa_success_requires_manual_visual_review():
    qa = load_module("qa")
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        (root / "chapters").mkdir()
        (root / "workspace/intermediate").mkdir(parents=True)
        (root / "workspace/output").mkdir(parents=True)
        (root / "workspace/input").mkdir(parents=True)
        (root / "chapters/1_intro.tex").write_text("绪论\n正文内容完整。\n", encoding="utf-8")
        (root / "Reference.bib").write_text("% no cited entries\n", encoding="utf-8")
        (root / "main.pdf").write_bytes(b"%PDF-1.7\n")
        (root / "workspace/output/report.md").write_text("# Build Report\n", encoding="utf-8")
        (root / "workspace/input/metadata.yaml").write_text(
            "english_content_decision: omit\n",
            encoding="utf-8",
        )
        (root / "workspace/output/build_result.json").write_text(
            json.dumps(
                {
                    "status": "passed",
                    "new_pdf": True,
                    "steps": [{"exit_code": 0} for _ in range(4)],
                    "flow_b_gate": {
                        "status": "passed",
                        "thesis_json_fingerprint": {"sha256": "thesis", "size_bytes": 1},
                        "source_docx_fingerprint": {"sha256": "docx", "size_bytes": 1},
                    },
                }
            ),
            encoding="utf-8",
        )
        (root / "workspace/intermediate/thesis.json").write_text(
            json.dumps(
                {
                    "structure": {
                        "chapters": [
                            {
                                "title": "绪论",
                                "file": "chapters/1_intro.tex",
                                "block_ids": [],
                            }
                        ]
                    },
                    "source_blocks": [],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        pdf_text = "目录\n摘要\n绪论\n正文内容完整。\n参考文献\n"
        with (
            patch.object(
                qa,
                "check_flow_b_gate",
                return_value={
                    "status": "passed",
                    "thesis_json_fingerprint": {"sha256": "thesis", "size_bytes": 1},
                    "source_docx_fingerprint": {"sha256": "docx", "size_bytes": 1},
                },
            ),
            patch.object(qa, "extract_text_with_pdftotext", return_value=pdf_text),
            patch.object(qa, "count_pdf_pages", return_value=8),
        ):
            result = qa.qa(root)

        assert result["status"] == "ready_for_manual_review"
        assert result["manual_review_required"] is True
        checks = {check["name"]: check for check in result["checks"]}
        assert checks["signal_abstract_en"]["status"] == "passed"
        assert "明确选择省略" in checks["signal_abstract_en"]["detail"]
        assert (root / "workspace/output/qa_report.md").exists()

        changed_gate = {
            "status": "passed",
            "thesis_json_fingerprint": {"sha256": "changed", "size_bytes": 2},
            "source_docx_fingerprint": {"sha256": "docx", "size_bytes": 1},
        }
        with (
            patch.object(qa, "check_flow_b_gate", return_value=changed_gate),
            patch.object(qa, "extract_text_with_pdftotext", return_value=pdf_text),
            patch.object(qa, "count_pdf_pages", return_value=8),
        ):
            changed_result = qa.qa(root)

        changed_checks = {check["name"]: check for check in changed_result["checks"]}
        assert changed_result["status"] == "failed"
        assert changed_checks["flow_b_gate_current"]["status"] == "passed"
        assert changed_checks["build_flow_b_binding"]["status"] == "failed"


if __name__ == "__main__":
    test_import_docx_preserves_superscript_runs()
    test_check_template_requires_new_fonts_directory_layout()
    test_check_template_guides_template_download_fallbacks()
    test_prepare_workspace_requires_consent_then_archives_old_outputs()
    test_metadata_parser_preserves_numeric_identifiers_with_leading_zeroes()
    test_import_docx_preserves_image_anchor_order()
    test_import_docx_preserves_repeated_uses_of_the_same_image_relationship()
    test_export_assets_does_not_mark_image_semantic_position_mapped()
    test_import_docx_reports_unsupported_features()
    test_import_and_prescan_expand_block_content_controls()
    test_export_assets_blocks_changed_source_docx()
    test_flow_b_gate_blocks_unconfirmed_unsupported_features()
    test_flow_b_gate_blocks_manual_figure_reference_numbers()
    test_flow_b_gate_blocks_broken_latex_label_refs()
    test_check_env_reports_missing_required_latex_packages()
    test_check_env_python_docx_hint_uses_short_timeout_and_mirror_fallback()
    test_check_env_reports_missing_latex_commands_as_structured_issues()
    test_check_env_qa_stage_reports_optional_tools_without_blocking()
    test_check_env_main_allows_needs_review_exit_code()
    test_check_env_distinguishes_missing_kpsewhich_from_tex_packages()
    test_check_env_blocks_unsupported_python_version()
    test_check_env_detects_broken_python_docx_import()
    test_check_env_rejects_unrelated_docx_module_without_document_api()
    test_prescan_reads_cover_table_metadata_without_report_style_default()
    test_prescan_and_import_accept_table_only_docx_metadata()
    test_render_basicinfo_blocks_missing_report_style()
    test_render_basicinfo_blocks_missing_required_cover_metadata()
    test_render_basicinfo_blocks_missing_subtitle_when_enabled()
    test_render_basicinfo_blocks_unapproved_generated_english()
    test_render_basicinfo_requires_missing_english_content_decision()
    test_render_basicinfo_requires_explicit_source_field_evidence()
    test_render_basicinfo_marks_only_verified_metadata_bindings_rendered()
    test_render_basicinfo_requires_an_explicit_destination_for_residual_text()
    test_render_basicinfo_accepts_report_style_candidate_synonyms()
    test_render_basicinfo_blocks_malformed_ledger_without_writing()
    test_qa_flags_bibtex_and_citation_lint_failures()
    test_qa_ignores_stale_chapters_not_referenced_by_current_mainbody()
    test_qa_reads_parenthesized_bibtex_keys()
    test_render_bib_preserves_existing_file_until_mapping_is_confirmed()
    test_render_bib_does_not_partially_write_when_any_entry_is_unresolved()
    test_render_chapters_preserves_superscript_and_heading_levels()
    test_render_chapters_outputs_labels_and_reference_rewrites()
    test_latex_escape_ascii_double_quotes_and_single_scan()
    test_render_chapters_blocks_prefix_path_escape()
    test_render_chapters_blocks_duplicate_chapter_targets_before_writing()
    test_render_chapters_requires_an_exported_existing_image_asset()
    test_render_chapters_blocks_discarded_block_in_structure()
    test_render_basicinfo_supports_thesis_title_abs()
    test_render_basicinfo_hides_hyperref_link_borders()
    test_build_xelatex_uses_noninteractive_error_flags()
    test_build_timeout_preserves_byte_output_without_type_error()
    test_build_success_writes_machine_and_human_reports()
    test_build_stops_before_archiving_when_flow_b_gate_is_blocked()
    test_diagnose_build_classifies_standard_latex_missing_file_messages()
    test_diagnose_build_writes_actionable_result()
    test_flow_b_gate_validates_ledger_identity_and_chapter_ownership()
    test_flow_b_gate_reports_malformed_source_block_without_crashing()
    test_flow_b_gate_rechecks_source_docx_fingerprint_before_flow_c()
    test_render_chapters_table_uses_fixed_font_without_resizebox()
    test_qa_flags_missing_superscript_rendering_and_resizebox()
    test_qa_flags_broken_latex_label_refs()
    test_qa_placeholder_scan_includes_generated_chapter_files()
    test_qa_requires_build_result_for_pdf_freshness()
    test_qa_reports_malformed_json_inputs_without_crashing()
    test_qa_unverifiable_page_count_requires_review_instead_of_false_failure()
    test_qa_counts_pages_with_pdfinfo_before_byte_fallback()
    test_qa_tool_timeouts_degrade_without_crashing()
    test_qa_body_signal_tolerates_malformed_ledger_sections()
    test_qa_success_requires_manual_visual_review()
    print("DOCX fidelity regression tests passed")
